
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import datetime

# Function to generate Word report
def generate_report(swot, pestel, ratios, competitors, company, industry):
    today = datetime.date.today().strftime("%B %d, %Y")
    doc = Document()
    doc.add_heading('Market Analysis Report', 0)
    doc.add_paragraph(f'Date: {today}')
    doc.add_paragraph(f'Company: {company}')
    doc.add_paragraph(f'Industry: {industry}')

    doc.add_heading('1. SWOT Analysis', level=1)
    for key, value in swot.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('2. PESTEL Analysis', level=1)
    for key, value in pestel.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('3. Financial Ratios', level=1)
    for key, value in ratios.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('4. Competitor Benchmarking', level=1)
    for _, row in competitors.iterrows():
        doc.add_paragraph(
            f"• {row['Competitor']}: Market Share - {row['Market Share (%)']}, "
            f"Strengths - {row['Strengths']}, Weaknesses - {row['Weaknesses']}, "
            f"Price - {row['Price Range (₦)']}, Quality - {row['Product Quality']}, "
            f"Channels - {row['Distribution Channels']}"
        )

    doc.add_paragraph("\n📄 Auto-generated using Streamlit Market Analysis Tool.")
    return doc

st.set_page_config(page_title="Market Analysis App", layout="wide")
st.title("📊 Market Analysis & Report Generator")

# Sidebar info
st.sidebar.header("Company Details")
company = st.sidebar.text_input("Company Name", "VoltEdge Energy Solutions")
industry = st.sidebar.text_input("Industry", "Storage Battery")

# SWOT input
st.header("1️⃣ SWOT Analysis")
swot_cols = st.columns(4)
swot_labels = ["Strengths", "Weaknesses", "Opportunities", "Threats"]
swot_data = {label: col.text_area(label, height=150) for col, label in zip(swot_cols, swot_labels)}

# PESTEL input
st.header("2️⃣ PESTEL Analysis")
pestel_labels = ["Political", "Economic", "Social", "Technological", "Environmental", "Legal"]
pestel_data = {label: st.text_area(f"{label} Factors", height=100) for label in pestel_labels}

# Financials input and chart
st.header("3️⃣ Financial Overview + Chart")
col1, col2 = st.columns(2)
with col1:
    revenue = st.number_input("Revenue (₦)", value=120_000_000)
    cogs = st.number_input("Cost of Goods Sold (₦)", value=70_000_000)
    net_profit = st.number_input("Net Profit (₦)", value=15_000_000)
    net_income = st.number_input("Net Income (₦)", value=15_000_000)
with col2:
    current_assets = st.number_input("Current Assets (₦)", value=30_000_000)
    current_liabilities = st.number_input("Current Liabilities (₦)", value=10_000_000)
    total_debt = st.number_input("Total Debt (₦)", value=20_000_000)
    equity = st.number_input("Shareholders' Equity (₦)", value=40_000_000)

# Financial chart
fig1, ax1 = plt.subplots()
labels = ['Revenue', 'COGS', 'Net Profit']
values = [revenue, cogs, net_profit]
bars = ax1.bar(labels, values, color=['green', 'red', 'blue'])
ax1.set_title("Financial Overview")
ax1.set_ylabel("Amount (₦)")
for bar in bars:
    yval = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width()/2.0, yval + 1e6, f"₦{int(yval/1e6)}M", ha='center', va='bottom')
st.pyplot(fig1)

# Financial ratios
ratios = {
    "Gross Margin (%)": f"{((revenue - cogs) / revenue) * 100:.2f}%" if revenue else "-",
    "Net Profit Margin (%)": f"{(net_profit / revenue) * 100:.2f}%" if revenue else "-",
    "Current Ratio": f"{(current_assets / current_liabilities):.2f}" if current_liabilities else "-",
    "Debt-to-Equity Ratio": f"{(total_debt / equity):.2f}" if equity else "-",
    "Return on Equity (ROE) (%)": f"{(net_income / equity) * 100:.2f}%" if equity else "-"
}

# Competitor input + chart
st.header("4️⃣ Competitor Benchmarking + Market Share Chart")
comp_df = pd.DataFrame({
    "Competitor": ["VoltEdge", "Competitor A", "Competitor B"],
    "Market Share (%)": [50, 30, 20],
    "Strengths": ["", "", ""],
    "Weaknesses": ["", "", ""],
    "Price Range (₦)": ["", "", ""],
    "Product Quality": ["", "", ""],
    "Distribution Channels": ["", "", ""]
})
comp_df = st.data_editor(comp_df, num_rows="dynamic")

# Pie chart
fig2, ax2 = plt.subplots()
ax2.pie(comp_df['Market Share (%)'],
        labels=comp_df['Competitor'],
        autopct='%1.1f%%',
        startangle=90)
ax2.axis('equal')
ax2.set_title("Competitor Market Share")
st.pyplot(fig2)

# Export report
if st.button("📤 Export Report to Word"):
    report = generate_report(swot_data, pestel_data, ratios, comp_df, company, industry)
    path = "Market_Analysis_Report.docx"
    report.save(path)
    with open(path, "rb") as f:
        st.download_button("📄 Download Market Analysis Report", f, file_name=path)
