import streamlit as st
import pandas as pd
from docx import Document
import datetime

# --- Function to create Word report ---
def generate_report(swot, pestel, ratios, competitors, company, industry):
    today = datetime.date.today().strftime("%B %d, %Y")
    doc = Document()
    doc.add_heading('Market Analysis Report', 0)
    doc.add_paragraph(f'Date: {today}')
    doc.add_paragraph(f'Company: {company}')
    doc.add_paragraph(f'Industry: {industry}')

    doc.add_heading('1. SWOT Analysis', level=1)
    for key, value in swot.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('2. PESTEL Analysis', level=1)
    for key, value in pestel.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('3. Financial Ratios', level=1)
    for key, value in ratios.items():
        doc.add_paragraph(f"• {key}: {value}")

    doc.add_heading('4. Competitor Benchmarking', level=1)
    for _, row in competitors.iterrows():
        line = f"• {row['Competitor']}: Market Share - {row['Market Share (%)']}, " \
               f"Strengths - {row['Strengths']}, Weaknesses - {row['Weaknesses']}, " \
               f"Price - {row['Price Range (₦)']}, Quality - {row['Product Quality']}, " \
               f"Channels - {row['Distribution Channels']}"
        doc.add_paragraph(line)

    doc.add_paragraph("\n📄 Auto-generated using Streamlit Market Analysis Tool.")
    return doc

# --- Streamlit UI ---
st.set_page_config(page_title="Market Analysis Tool", layout="wide")
st.title("📊 Market Analysis Automation Tool")

# Sidebar company info
st.sidebar.header("🔍 Company Information")
industry = st.sidebar.text_input("Industry", "Storage Battery")
company = st.sidebar.text_input("Company Name", "VoltEdge Energy Solutions")

# SWOT
st.header("1. SWOT Analysis")
swot_cols = st.columns(4)
swot_labels = ["Strengths", "Weaknesses", "Opportunities", "Threats"]
swot_data = {}
for col, label in zip(swot_cols, swot_labels):
    swot_data[label] = col.text_area(label, height=150)

# PESTEL
st.header("2. PESTEL Analysis")
pestel_labels = ["Political", "Economic", "Social", "Technological", "Environmental", "Legal"]
pestel_data = {}
for label in pestel_labels:
    pestel_data[label] = st.text_area(f"{label} Factors", height=100)

# Financials
st.header("3. Financial Inputs")
col1, col2 = st.columns(2)
with col1:
    revenue = st.number_input("Revenue (₦)", 0.0)
    cogs = st.number_input("Cost of Goods Sold (₦)", 0.0)
    net_profit = st.number_input("Net Profit (₦)", 0.0)
    net_income = st.number_input("Net Income (₦)", 0.0)
with col2:
    current_assets = st.number_input("Current Assets (₦)", 0.0)
    current_liabilities = st.number_input("Current Liabilities (₦)", 1.0)
    total_debt = st.number_input("Total Debt (₦)", 0.0)
    equity = st.number_input("Shareholders' Equity (₦)", 1.0)

# Calculate financial ratios
ratios = {
    "Gross Margin (%)": f"{((revenue - cogs) / revenue) * 100:.2f}%" if revenue else "-",
    "Net Profit Margin (%)": f"{(net_profit / revenue) * 100:.2f}%" if revenue else "-",
    "Current Ratio": f"{(current_assets / current_liabilities):.2f}" if current_liabilities else "-",
    "Debt-to-Equity Ratio": f"{(total_debt / equity):.2f}" if equity else "-",
    "Return on Equity (ROE) (%)": f"{(net_income / equity) * 100:.2f}%" if equity else "-"
}

# Competitor table
st.header("4. Competitor Benchmarking")
st.markdown("Enter details of your competitors:")
competitor_df = pd.DataFrame({
    "Competitor": ["Competitor A", "Competitor B"],
    "Market Share (%)": ["", ""],
    "Strengths": ["", ""],
    "Weaknesses": ["", ""],
    "Price Range (₦)": ["", ""],
    "Product Quality": ["", ""],
    "Distribution Channels": ["", ""]
})
edited_competitor_df = st.data_editor(competitor_df, num_rows="dynamic")

# Export to Word
if st.button("📤 Export Report to Word"):
    report = generate_report(swot_data, pestel_data, ratios, edited_competitor_df, company, industry)
    docx_path = "Market_Analysis_Report.docx"
    report.save(docx_path)
    with open(docx_path, "rb") as file:
        st.download_button("📄 Download Market Analysis Report", file, file_name=docx_path)

st.success("✅ Entered data is ready for export. Fill in all fields and click 'Export Report'.")
